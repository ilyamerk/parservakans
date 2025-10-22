# build_report_docx.py
import argparse, pandas as pd, numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

COLS = ["Должность","Работодатель","ЗП от (т.р.)","ЗП до (т.р.)","В час",
        "Требуемый\nопыт","Труд-во","График","Частота \nвыплат","Льготы","Ссылка"]

def load_df(p: Path) -> pd.DataFrame:
    df = pd.read_csv(p)
    for c in COLS:
        if c not in df.columns:
            df[c] = np.nan
    # приведение типов
    for c in ["ЗП от (т.р.)","ЗП до (т.р.)","В час"]:
        df[c] = pd.to_numeric(df[c], errors="coerce")
    return df[COLS]

def freq_series(s: pd.Series, top=8):
    vc = (s.dropna().astype(str).str.lower()
          .str.replace(r"\s+", " ", regex=True)
          .str.split(r"[;,]", regex=True).explode().str.strip())
    vc = vc[vc.ne("")].value_counts()
    return vc.head(top)

def fmt(v, suf=""):
    if v is None or (isinstance(v, float) and np.isnan(v)): return "—"
    if isinstance(v, float): return f"{v:.0f}{suf}"
    return f"{v}{suf}"

def add_heading(doc, text, lvl=1):
    h = doc.add_heading(text, level=lvl)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_kv(doc, k, v):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{k}: "); r1.bold = True
    p.add_run(str(v))

def add_list(doc, title, series, total_n: int):
    add_heading(doc, title, 2)
    if series is None or series.empty or total_n == 0:
        doc.add_paragraph("—")
        return
    for k, v in series.items():
        pct = (int(v) / total_n) * 100.0
        doc.add_paragraph(f"{k}: {int(v)} ({pct:.0f}%)", style="List Bullet")

def add_top_table(doc, df_top):
    add_heading(doc, "ТОП-10 по ставке «В час»", 2)
    if df_top.empty:
        doc.add_paragraph("—")
        return
    t = doc.add_table(rows=1, cols=5)
    hdr = t.rows[0].cells
    hdr[0].text = "Должность"
    hdr[1].text = "Работодатель"
    hdr[2].text = "В час, ₽"
    hdr[3].text = "ЗП (т.р.)"
    hdr[4].text = "Ссылка"
    for _, r in df_top.iterrows():
        row = t.add_row().cells
        row[0].text = str(r["Должность"] or "")
        row[1].text = str(r["Работодатель"] or "")
        row[2].text = fmt(r["В час"], "")
        mid = np.nan
        if pd.notna(r["ЗП от (т.р.)"]) and pd.notna(r["ЗП до (т.р.)"]):
            mid = (r["ЗП от (т.р.)"] + r["ЗП до (т.р.)"]) / 2
        elif pd.notna(r["ЗП от (т.р.)"]):
            mid = r["ЗП от (т.р.)"]
        elif pd.notna(r["ЗП до (т.р.)"]):
            mid = r["ЗП до (т.р.)"]
        row[3].text = fmt(mid, "")
        row[4].text = str(r["Ссылка"] or "")

def save_docx_safely(doc: Document, out_path: str, retries: int = 6, delay: float = 1.0):
    """Безопасное сохранение с заменой занятого файла."""
    import os, time
    out = Path(out_path)
    tmp = out.with_suffix(out.suffix + ".tmp")
    doc.save(tmp)
    for _ in range(retries):
        try:
            os.replace(tmp, out)
            print("OK:", out)
            return
        except PermissionError:
            time.sleep(delay)
    stamped = out.with_name(out.stem + f"_{int(time.time())}" + out.suffix)
    os.replace(tmp, stamped)
    print("Файл занят. Сохранил как:", stamped)

def main():
    ap = argparse.ArgumentParser(description="DOCX-отчёт по вакансиям")
    ap.add_argument("--input_csv", required=True)        # parsers/raw.csv
    ap.add_argument("--output_docx", required=True)      # Отчёт_....docx
    ap.add_argument("--query", default="")
    ap.add_argument("--city", default="")
    a = ap.parse_args()

    df = load_df(Path(a.input_csv))
    n = int(len(df))

    # ===== очистка чисел =====
    # почасовая: только положительные и разумные
    df["В час"] = pd.to_numeric(df["В час"], errors="coerce")
    df.loc[(df["В час"].isna()) | (df["В час"] <= 0), "В час"] = np.nan
    vh = df["В час"].dropna()
    vh = vh[(vh > 50) & (vh < 100000)]

    # месячная: берем середину вилки, если оба края есть; иначе имеющийся край
    for c in ["ЗП от (т.р.)","ЗП до (т.р.)"]:
        df[c] = pd.to_numeric(df[c], errors="coerce")
        df.loc[(df[c].isna()) | (df[c] <= 0), c] = np.nan
    zpf, zpt = df["ЗП от (т.р.)"], df["ЗП до (т.р.)"]
    mid = ((zpf + zpt) / 2).where(zpf.notna() & zpt.notna())
    monthly = mid.fillna(zpf).fillna(zpt)
    monthly = monthly[(monthly.notna()) & (monthly > 10) & (monthly < 10000)]

    # метрики
    vh_med = float(vh.median()) if not vh.empty else None
    vh_min = float(vh.min()) if not vh.empty else None
    vh_max = float(vh.max()) if not vh.empty else None

    m_med = float(monthly.median()) if not monthly.empty else None
    m_min = float(monthly.min()) if not monthly.empty else None
    m_max = float(monthly.max()) if not monthly.empty else None

    # частоты
    top_benefits = freq_series(df["Льготы"])
    exp_freq     = df["Требуемый\nопыт"].dropna().astype(str).value_counts().head(6)
    empl_freq    = df["Труд-во"].dropna().astype(str).value_counts().head(6)
    sched_fr     = freq_series(df["График"])
    pay_fr       = df["Частота \nвыплат"].dropna().astype(str).value_counts().head(6)

    # топ-10 по «В час»
    top10 = df.dropna(subset=["В час"]).sort_values("В час", ascending=False).head(10)

    # ===== DOCX =====
    doc = Document()
    title = f"Аналитическая записка: «{a.query}», {a.city}".strip(", ")
    add_heading(doc, title, 0)

    add_heading(doc, "Сводные метрики", 1)
    add_kv(doc, "Вакансий", n)
    add_kv(doc, "«В час», медиана", fmt(vh_med, " ₽"))
    add_kv(doc, "мин", fmt(vh_min, " ₽"))
    add_kv(doc, "макс", fmt(vh_max, " ₽"))
    add_kv(doc, "ЗП (т.р.), медиана", fmt(m_med, ""))
    add_kv(doc, "мин", fmt(m_min, ""))
    add_kv(doc, "макс", fmt(m_max, ""))

    add_list(doc, "Льготы (ТОП)", top_benefits, n)
    add_list(doc, "Требуемый опыт", exp_freq, n)
    add_list(doc, "Тип оформления (ТК/ГПХ)", empl_freq, n)
    add_list(doc, "Графики", sched_fr, n)
    add_list(doc, "Частота выплат", pay_fr, n)

    add_top_table(doc, top10)

    # базовый шрифт
    for s in doc.styles:
        if s.type == 1 and s.font.name is None:
            s.font.name = "Calibri"
            s.font.size = Pt(11)

    Path(a.output_docx).parent.mkdir(parents=True, exist_ok=True)
    save_docx_safely(doc, a.output_docx)

if __name__ == "__main__":
    main()
